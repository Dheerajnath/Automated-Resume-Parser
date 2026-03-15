import pdfplumber
from docx import Document


def extract_text_from_pdf(filepath: str) -> str:
    """
    Extract all text from a PDF file using pdfplumber.
    Returns concatenated text from all pages.
    """
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    return text.strip()


def extract_text_from_docx(filepath: str) -> str:
    """
    Extract all text from a DOCX file using python-docx.
    Returns concatenated paragraphs.
    """
    text = ""
    try:
        doc = Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Also extract text from tables within the DOCX
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    return text.strip()


def extract_text(filepath: str, file_extension: str) -> str:
    """
    Dispatch to the correct extractor based on file extension.
    Supported: .pdf, .docx
    """
    ext = file_extension.lower().lstrip(".")
    if ext == "pdf":
        return extract_text_from_pdf(filepath)
    elif ext == "docx":
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
